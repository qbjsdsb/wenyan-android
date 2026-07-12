"""pipeline_runner.py - 文研App资料数字化处理管线

manifest驱动的断点续传批处理脚本，将wenyanziliao文件夹的208个文件
数字化为结构化JSON，作为App种子数据。

核心功能：
  1. 读取manifest.json，找出所有status="pending"的文件
  2. 按file_type路由到对应处理函数（PDF/DOCX/DOC/XLSX/XLS/IMAGE/ZIP）
  3. 指数退避重试（最多3次失败后跳过，标记status="failed"）
  4. 每处理完一个文件立即更新manifest.json（断点续传）
  5. 输出结果到output目录（每个文件一个JSON）

对应 Task 1.3（批处理骨架）+ Task 1.5（非PDF文件处理路由）。

处理路由：
  - pdf/native    → pdfplumber直接提取文本
  - pdf/ocr_layer → pdfplumber提取现有OCR文本层
  - pdf/scan_only → MinerU CLI (mineru -p <input> -o <output> -m auto -b pipeline)
  - pdf/mixed     → 文本层页pdfplumber + 扫描页MinerU
  - docx          → python-docx提取文本和表格
  - doc           → pywin32(win32com)转.docx后提取（需MS Word，USERPROFILE需指向C盘）
  - xlsx          → openpyxl提取表格数据
  - xls           → xlrd 2.0+提取表格数据
  - image         → MinerU CLI pipeline后端OCR（原PaddleOCR，2026-07-10因3.x不兼容改用MinerU）
  - zip           → 解压后递归处理内部文件
"""

import argparse
import json
import os
import shutil
import subprocess
import sys
import tempfile
import threading
import time
from concurrent.futures import ThreadPoolExecutor, as_completed
from datetime import datetime

# ===== 配置D盘环境（必须在其他导入之前，避免C盘写入） =====
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import d_drive_env  # noqa: F401 - 导入即自动配置所有环境变量到D盘


# ===== 常量定义 =====

# 最大重试次数（3次失败后标记为failed）
MAX_RETRIES = 3

# 指数退避基础延迟（秒），实际延迟 = BASE_BACKOFF * 2^(attempt-1)
BASE_BACKOFF_SECONDS = 2

# MinerU CLI命令模板（完整路径，避免PATH查找失败）
MINERU_CMD = r"C:\Users\33425\miniconda3\envs\ocr\Scripts\mineru.exe"

# 内容来源类型
SOURCE_NATIVE = "TEXTBOOK_NATIVE"
SOURCE_OCR = "TEXTBOOK_OCR"

# OCR状态
OCR_VERIFIED = "VERIFIED"
OCR_PENDING = "PENDING"

# 快速文件并行处理的工作线程数（MinerU慢速文件串行，不受此限制）
DEFAULT_FAST_WORKERS = 4


def get_file_priority(file_record):
    """获取文件处理优先级（数字越小越优先处理）。

    快速文件优先（秒级完成），慢速OCR文件后处理。
    同优先级内按页数升序（小文件先处理，快速看到进展）。

    注意：doc和zip归类为慢速，因为：
    - doc需要Word COM，在线程池中COM初始化不兼容（RPC_E_CALL_REJECTED）
    - zip可能内部含scan_only/mixed PDF，递归处理时需MinerU OCR

    Returns:
        int: 优先级（1最快 ~ 6最慢）
    """
    file_type = file_record.get("file_type", "")
    pdf_type = file_record.get("pdf_type", "")

    if file_type in ("docx", "xlsx", "xls"):
        return 1  # 最快（秒级，线程安全）
    elif file_type == "pdf" and pdf_type in ("native", "ocr_layer"):
        return 2  # 快（pdfplumber秒级）
    elif file_type == "pdf" and pdf_type == "mixed":
        return 3  # 中等（优化后只OCR扫描页）
    elif file_type in ("image", "doc", "zip"):
        return 4  # 慢（doc需Word COM串行；zip可能含OCR PDF；image需MinerU）
    elif file_type == "pdf" and pdf_type == "scan_only":
        return 5  # 最慢（MinerU全量OCR）
    else:
        return 6  # 未知类型最后


# ===== PDF处理函数 =====

def process_pdf_native(path):
    """用pdfplumber提取原生电子文本PDF。

    对应Spec：NATIVE类型PDF（28个），零错字直接提取。
    content_source = 'TEXTBOOK_NATIVE', ocr_status = 'VERIFIED'

    Args:
        path: PDF文件绝对路径

    Returns:
        dict: 提取结果，包含pages和tables
    """
    import pdfplumber

    pages = []
    tables = []

    with pdfplumber.open(path) as pdf:
        for i, page in enumerate(pdf.pages):
            text = page.extract_text() or ""
            pages.append({
                "page_num": i + 1,
                "text": text.strip(),
                "char_count": len(text.strip()),
            })
            # 提取表格
            page_tables = page.extract_tables()
            for table_idx, table in enumerate(page_tables):
                tables.append({
                    "page_num": i + 1,
                    "table_idx": table_idx,
                    "rows": table,
                })

    return {
        "pages": pages,
        "tables": tables,
        "total_pages": len(pages),
        "content_source": SOURCE_NATIVE,
        "ocr_status": OCR_VERIFIED,
    }


def process_pdf_ocr_layer(path):
    """用pdfplumber提取OCR文本层PDF。

    对应Spec：OCR_LAYER类型PDF（6个），提取现有OCR文本层后需抽样校对。
    content_source = 'TEXTBOOK_OCR', ocr_status = 'PENDING'

    Args:
        path: PDF文件绝对路径

    Returns:
        dict: 提取结果，包含pages
    """
    import pdfplumber

    pages = []

    with pdfplumber.open(path) as pdf:
        for i, page in enumerate(pdf.pages):
            text = page.extract_text() or ""
            pages.append({
                "page_num": i + 1,
                "text": text.strip(),
                "char_count": len(text.strip()),
            })

    return {
        "pages": pages,
        "total_pages": len(pages),
        "content_source": SOURCE_OCR,
        "ocr_status": OCR_PENDING,
    }


def process_pdf_scan_only(path, work_dir):
    """用MinerU CLI处理扫描件PDF。

    对应Spec：SCAN_ONLY类型PDF（111个），MinerU 3.x OCR。
    CLI命令：mineru -p <input> -o <output> -m auto
    content_source = 'TEXTBOOK_OCR', ocr_status = 'PENDING'

    Args:
        path: PDF文件绝对路径
        work_dir: 临时工作目录（MinerU输出目录）

    Returns:
        dict: 提取结果，包含pages和MinerU输出路径
    """
    file_name = os.path.basename(path)
    mineru_output = os.path.join(work_dir, "mineru_output")

    # 调用MinerU CLI：mineru -p <input> -o <output> -m auto -b pipeline -f False
    # -b pipeline：使用pipeline后端（CPU可用，无需GPU；hybrid-engine默认需要CUDA）
    # -f False：禁用公式解析（文学资料无需公式识别，且Unimernet模型加载失败会阻断处理）
    cmd = [MINERU_CMD, "-p", path, "-o", mineru_output, "-m", "auto", "-b", "pipeline", "-f", "False"]
    result = subprocess.run(
        cmd,
        capture_output=True,
        text=True,
        timeout=7200,  # 2小时超时（大文件可能很慢）
    )

    if result.returncode != 0:
        raise RuntimeError(
            f"MinerU处理失败 (返回码={result.returncode}): "
            f"{result.stderr[:500]}"
        )

    # 查找MinerU输出的content_list.json
    content_list_path = _find_mineru_output(mineru_output, file_name)
    if content_list_path is None:
        raise RuntimeError(
            f"MinerU输出未找到content_list.json，输出目录: {mineru_output}"
        )

    # 解析content_list.json
    with open(content_list_path, "r", encoding="utf-8") as f:
        content_list = json.load(f)

    # 将content_list转换为pages格式
    # MinerU v1格式：[{type, text/table_body}, ...]（扁平列表）
    # MinerU v2格式：[ [{type, content}, ...], ...]（按页分组嵌套）
    pages = []
    import re as _re
    import html as _html

    def _extract_text_from_item(item):
        """从content_list项提取纯文本。"""
        if isinstance(item, str):
            return item
        if not isinstance(item, dict):
            return ""
        item_type = item.get("type", "")
        # v2格式的文本在content字段内
        if "content" in item and isinstance(item["content"], dict):
            content = item["content"]
            if item_type == "table":
                html_str = content.get("html", "")
            else:
                html_str = content.get("text", "") or content.get("html", "")
        else:
            # v1格式
            if item_type == "table":
                html_str = item.get("table_body", "")
            else:
                html_str = item.get("text", "")
        # 如果是HTML表格，提取单元格文本
        if html_str and "<td" in html_str:
            # 简单HTML转文本：替换</td>为制表符，</tr>为换行，去标签
            text = _re.sub(r"<td[^>]*>", "", html_str)
            text = text.replace("</td>", "\t")
            text = _re.sub(r"<tr[^>]*>", "", text)
            text = text.replace("</tr>", "\n")
            text = _re.sub(r"<[^>]+>", "", text)
            text = _html.unescape(text)
            return text.strip()
        return (html_str or "").strip()

    # 判断是v1（扁平）还是v2（嵌套）格式
    if content_list and isinstance(content_list[0], list):
        # v2格式：按页分组
        for i, page_items in enumerate(content_list):
            page_text = "\n".join(_extract_text_from_item(item) for item in page_items)
            pages.append({
                "page_num": i + 1,
                "text": page_text.strip(),
                "char_count": len(page_text.strip()),
            })
    else:
        # v1格式：扁平列表（每项代表一个内容块，不严格对应页）
        all_text = []
        for item in content_list:
            text = _extract_text_from_item(item)
            if text:
                all_text.append(text)
        # v1扁平列表无页码信息，合并为一页
        combined_text = "\n".join(all_text)
        pages.append({
            "page_num": 1,
            "text": combined_text,
            "char_count": len(combined_text),
        })

    return {
        "pages": pages,
        "total_pages": len(pages),
        "mineru_output_dir": mineru_output,
        "content_source": SOURCE_OCR,
        "ocr_status": OCR_PENDING,
    }


def process_pdf_mixed(path, work_dir):
    """处理混合类型PDF（部分页有文本层，部分页为扫描）。

    对应Spec：MIXED类型PDF（39个），文本层页用pdfplumber，扫描页用MinerU。
    content_source = 'TEXTBOOK_OCR', ocr_status = 'PENDING'

    优化（2026-07-11）：只提取扫描页生成临时PDF，仅对扫描页调用MinerU，
    避免对已有文本层的页面重复OCR，大幅减少MinerU处理时间。

    Args:
        path: PDF文件绝对路径
        work_dir: 临时工作目录

    Returns:
        dict: 提取结果，包含pages（混合提取）
    """
    import pdfplumber
    import fitz  # PyMuPDF，用于提取扫描页生成临时PDF

    # 第一步：用pdfplumber逐页提取，识别有文本和无文本的页面
    pdfplumber_pages = []
    scan_page_indices = []

    with pdfplumber.open(path) as pdf:
        for i, page in enumerate(pdf.pages):
            text = (page.extract_text() or "").strip()
            pdfplumber_pages.append(text)
            if len(text) == 0:
                scan_page_indices.append(i)

    # 第二步：如果有扫描页，只提取扫描页生成临时PDF，仅对临时PDF调用MinerU
    mineru_pages = {}  # {原PDF页码(1-based): 文本}
    if scan_page_indices:
        # 用PyMuPDF提取扫描页生成临时PDF（比整个PDF送MinerU快得多）
        src_doc = fitz.open(path)
        scan_doc = fitz.open()  # 新空PDF
        for idx in scan_page_indices:
            scan_doc.insert_pdf(src_doc, from_page=idx, to_page=idx)
        src_doc.close()

        temp_pdf_path = os.path.join(work_dir, "scan_pages_only.pdf")
        scan_doc.save(temp_pdf_path)
        scan_doc.close()

        print(f"  mixed优化：原PDF {len(pdfplumber_pages)} 页，"
              f"扫描页 {len(scan_page_indices)} 页，"
              f"仅对扫描页调用MinerU（省{len(pdfplumber_pages) - len(scan_page_indices)}页OCR）")

        # 只对扫描页临时PDF调用MinerU
        mineru_result = process_pdf_scan_only(temp_pdf_path, work_dir)
        # MinerU输出的页码是临时PDF的页码（1到N），映射回原PDF页码
        mineru_result_pages = mineru_result.get("pages", [])
        for i, page in enumerate(mineru_result_pages):
            if i < len(scan_page_indices):
                original_page_num = scan_page_indices[i] + 1  # 0-based转1-based
                mineru_pages[original_page_num] = page["text"]

    # 第三步：合并结果——文本层页用pdfplumber，扫描页用MinerU
    pages = []
    for i, text in enumerate(pdfplumber_pages):
        page_num = i + 1
        if len(text) > 0:
            # 有文本层，用pdfplumber结果
            pages.append({
                "page_num": page_num,
                "text": text,
                "char_count": len(text),
                "source": "pdfplumber",
            })
        else:
            # 无文本层（扫描页），用MinerU结果
            mineru_text = mineru_pages.get(page_num, "")
            pages.append({
                "page_num": page_num,
                "text": mineru_text,
                "char_count": len(mineru_text),
                "source": "mineru",
            })

    return {
        "pages": pages,
        "total_pages": len(pages),
        "scan_page_count": len(scan_page_indices),
        "content_source": SOURCE_OCR,
        "ocr_status": OCR_PENDING,
    }


def _find_mineru_output(output_dir, input_filename):
    """查找MinerU输出的content_list.json文件路径。

    MinerU 3.x的实际输出结构：
      output_dir/<stem>/auto/<stem>_content_list.json
      output_dir/<stem>/auto/<stem>_content_list_v2.json

    也兼容旧版本的可能结构。

    Args:
        output_dir: MinerU输出根目录
        input_filename: 输入文件名

    Returns:
        content_list.json的完整路径，如果未找到返回None
    """
    stem = os.path.splitext(input_filename)[0]
    # MinerU 3.x实际输出路径（带auto子目录和文件名前缀）
    possible_paths = [
        os.path.join(output_dir, stem, "auto", f"{stem}_content_list.json"),
        os.path.join(output_dir, stem, "auto", f"{stem}_content_list_v2.json"),
        os.path.join(output_dir, stem, "auto", "content_list.json"),
        os.path.join(output_dir, stem, "content_list.json"),
        os.path.join(output_dir, stem, f"{stem}_content_list.json"),
        os.path.join(output_dir, "content_list.json"),
    ]
    for path in possible_paths:
        if os.path.exists(path):
            return path
    # 递归搜索任何包含content_list的json文件
    for root, dirs, files in os.walk(output_dir):
        for file in files:
            if "content_list" in file and file.endswith(".json"):
                return os.path.join(root, file)
    return None


# ===== DOCX/DOC处理函数 =====

def process_docx(path):
    """用python-docx提取DOCX文本和表格。

    对应Spec：DOCX（31个），python-docx提取（仅支持.docx）。
    content_source = 'TEXTBOOK_NATIVE', ocr_status = 'VERIFIED'

    Args:
        path: DOCX文件绝对路径

    Returns:
        dict: 提取结果，包含paragraphs和tables
    """
    from docx import Document

    doc = Document(path)

    # 提取段落文本
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # 提取表格
    tables = []
    for table_idx, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)
        tables.append({
            "table_idx": table_idx,
            "rows": rows,
        })

    return {
        "paragraphs": paragraphs,
        "tables": tables,
        "total_paragraphs": len(paragraphs),
        "content_source": SOURCE_NATIVE,
        "ocr_status": OCR_VERIFIED,
    }


def process_doc(path):
    """用pywin32(win32com)转换DOC为DOCX后提取。

    对应Spec：DOC（3个），pywin32调用MS Word COM接口转换。
    python-docx不支持.doc，必须先转换为.docx。
    需要本机安装MS Word。

    注意：d_drive_env.py重定向了USERPROFILE/HOME到D盘，会导致Word COM
    无法找到模板路径，Documents.Open返回None。因此在调用Word COM前必须
    临时恢复C盘原始环境变量。

    Args:
        path: DOC文件绝对路径

    Returns:
        dict: 提取结果（与process_docx格式一致）

    Raises:
        RuntimeError: MS Word不可用或转换失败
    """
    import win32com.client
    import pythoncom

    # 保存d_drive_env设置的环境变量，临时恢复C盘原始值
    # Word COM依赖USERPROFILE/APPDATA定位模板路径，重定向后Open返回None
    C_USERPROFILE = r"C:\Users\33425"
    saved_env = {}
    for key in ["USERPROFILE", "HOME", "HOMEPATH"]:
        saved_env[key] = os.environ.get(key)
        os.environ[key] = C_USERPROFILE
    # APPDATA/LOCALAPPDATA可能未被d_drive_env修改，但确保指向C盘
    saved_env["APPDATA"] = os.environ.get("APPDATA")
    os.environ["APPDATA"] = os.path.join(C_USERPROFILE, "AppData", "Roaming")
    saved_env["LOCALAPPDATA"] = os.environ.get("LOCALAPPDATA")
    os.environ["LOCALAPPDATA"] = os.path.join(C_USERPROFILE, "AppData", "Local")

    # 创建临时docx文件路径
    temp_dir = tempfile.mkdtemp(prefix="wenyan_doc_")
    docx_path = os.path.join(
        temp_dir,
        os.path.splitext(os.path.basename(path))[0] + ".docx"
    )

    word_app = None
    doc = None
    try:
        pythoncom.CoInitialize()
        # 启动Word COM对象
        word_app = win32com.client.Dispatch("Word.Application")
        word_app.Visible = False
        word_app.DisplayAlerts = 0  # wdAlertsNone

        # 打开DOC文件
        doc = word_app.Documents.Open(os.path.abspath(path))

        if doc is None:
            raise RuntimeError(
                "Word Documents.Open返回None（可能模板路径异常）"
            )

        # 另存为DOCX格式（FileFormat=16 = wdFormatXMLDocument）
        doc.SaveAs(os.path.abspath(docx_path), FileFormat=16)

        # 大文件（>5MB）的Close可能因Word内部超时而RPC_E_DISCONNECTED
        # SaveAs已完成，DOCX文件已写入磁盘，Close失败不影响数据
        try:
            doc.Close(SaveChanges=0)  # wdDoNotSaveChanges=0
        except Exception:
            pass  # 忽略Close错误，DOCX已保存
        doc = None

        # 验证DOCX文件已生成
        if not os.path.exists(docx_path):
            raise RuntimeError(
                f"DOCX文件未生成: {docx_path}"
            )

        # 用process_docx提取转换后的文件
        result = process_docx(docx_path)
        result["converted_from"] = "doc"
        return result

    except Exception as e:
        raise RuntimeError(
            f"DOC转换失败（请确认已安装MS Word）: {e}"
        )
    finally:
        # 清理COM对象
        if doc is not None:
            try:
                doc.Close()
            except Exception:
                pass
        if word_app is not None:
            try:
                word_app.Quit()
            except Exception:
                pass
        try:
            pythoncom.CoUninitialize()
        except Exception:
            pass
        # 恢复d_drive_env的环境变量
        for key, val in saved_env.items():
            if val is not None:
                os.environ[key] = val
        # 清理临时文件
        try:
            shutil.rmtree(temp_dir, ignore_errors=True)
        except Exception:
            pass


# ===== 表格处理函数 =====

def process_xlsx(path):
    """用openpyxl提取XLSX表格数据。

    对应Spec：XLSX（3个），openpyxl提取（仅支持.xlsx）。
    content_source = 'TEXTBOOK_NATIVE', ocr_status = 'VERIFIED'

    Args:
        path: XLSX文件绝对路径

    Returns:
        dict: 提取结果，包含sheets
    """
    from openpyxl import load_workbook

    wb = load_workbook(path, data_only=True)
    sheets = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = []
        for row in ws.iter_rows(values_only=True):
            # 将None转为空字符串，保留其他值
            cells = [str(cell) if cell is not None else "" for cell in row]
            rows.append(cells)
        sheets.append({
            "sheet_name": sheet_name,
            "rows": rows,
            "row_count": len(rows),
        })

    wb.close()

    return {
        "sheets": sheets,
        "total_sheets": len(sheets),
        "content_source": SOURCE_NATIVE,
        "ocr_status": OCR_VERIFIED,
    }


def process_xls(path):
    """用xlrd 2.0+提取XLS表格数据。

    对应Spec：XLS（13个），xlrd 2.0+提取（支持.xls，2.0+移除了.xlsx支持）。
    content_source = 'TEXTBOOK_NATIVE', ocr_status = 'VERIFIED'

    Args:
        path: XLS文件绝对路径

    Returns:
        dict: 提取结果，包含sheets
    """
    import xlrd

    # xlrd 2.0+仅支持.xls格式，formatting_info在某些版本不支持
    wb = xlrd.open_workbook(path)
    sheets = []

    for sheet in wb.sheets():
        rows = []
        for row_idx in range(sheet.nrows):
            cells = []
            for col_idx in range(sheet.ncols):
                cell_value = sheet.cell_value(row_idx, col_idx)
                # 处理不同类型的单元格值
                if isinstance(cell_value, float) and cell_value == int(cell_value):
                    cells.append(str(int(cell_value)))
                elif cell_value == "":
                    cells.append("")
                else:
                    cells.append(str(cell_value))
            rows.append(cells)
        sheets.append({
            "sheet_name": sheet.name,
            "rows": rows,
            "row_count": len(rows),
        })

    return {
        "sheets": sheets,
        "total_sheets": len(sheets),
        "content_source": SOURCE_NATIVE,
        "ocr_status": OCR_VERIFIED,
    }


# ===== 图片处理函数 =====

def process_image(path, work_dir=None):
    """用MinerU CLI处理图片OCR。

    对应Spec：图片（4个JPG/PNG真题照片）。
    原设计用PaddleOCR，但paddlepaddle 3.x与pytorch存在DLL冲突且
    PaddleOCR 3.x API不兼容，改为用MinerU CLI处理（与PDF scan_only共用逻辑）。
    content_source = 'TEXTBOOK_OCR', ocr_status = 'PENDING'

    Args:
        path: 图片文件绝对路径
        work_dir: 临时工作目录（可选，MinerU输出目录）

    Returns:
        dict: 提取结果，包含text和pages
    """
    if work_dir is None:
        work_dir = tempfile.mkdtemp(prefix="wenyan_image_")

    file_name = os.path.basename(path)
    mineru_output = os.path.join(work_dir, "mineru_output")

    # 调用MinerU CLI处理图片（与PDF scan_only相同的命令）
    # -b pipeline：使用pipeline后端（CPU可用，传统OCR快速可靠）
    # -f False：禁用公式解析（文学资料无需公式识别，且Unimernet模型加载失败会阻断处理）
    # 不用默认hybrid后端（VLM/Qwen2VL在CPU上极慢，12分钟无输出）
    cmd = [MINERU_CMD, "-p", path, "-o", mineru_output, "-m", "auto", "-b", "pipeline", "-f", "False"]
    result = subprocess.run(
        cmd,
        capture_output=True,
        text=True,
        timeout=600,  # 10分钟超时（pipeline后端处理单张图片足够）
    )

    if result.returncode != 0:
        raise RuntimeError(
            f"MinerU处理图片失败 (返回码={result.returncode}): "
            f"{result.stderr[:500]}"
        )

    # 查找MinerU输出的content_list.json
    content_list_path = _find_mineru_output(mineru_output, file_name)
    if content_list_path is None:
        raise RuntimeError(
            f"MinerU输出未找到content_list.json，输出目录: {mineru_output}"
        )

    # 解析content_list.json（复用PDF scan_only的解析逻辑）
    with open(content_list_path, "r", encoding="utf-8") as f:
        content_list = json.load(f)

    import re as _re
    import html as _html

    def _extract_text_from_item(item):
        """从content_list项提取纯文本。"""
        if isinstance(item, str):
            return item
        if not isinstance(item, dict):
            return ""
        item_type = item.get("type", "")
        if "content" in item and isinstance(item["content"], dict):
            content = item["content"]
            if item_type == "table":
                html_str = content.get("html", "")
            else:
                html_str = content.get("text", "") or content.get("html", "")
        else:
            if item_type == "table":
                html_str = item.get("table_body", "")
            else:
                html_str = item.get("text", "")
        if html_str and "<td" in html_str:
            text = _re.sub(r"<td[^>]*>", "", html_str)
            text = text.replace("</td>", "\t")
            text = _re.sub(r"<tr[^>]*>", "", text)
            text = text.replace("</tr>", "\n")
            text = _re.sub(r"<[^>]+>", "", text)
            text = _html.unescape(text)
            return text.strip()
        return (html_str or "").strip()

    # 判断是v1（扁平）还是v2（嵌套）格式
    pages = []
    if content_list and isinstance(content_list[0], list):
        for i, page_items in enumerate(content_list):
            page_text = "\n".join(_extract_text_from_item(item) for item in page_items)
            pages.append({
                "page_num": i + 1,
                "text": page_text.strip(),
                "char_count": len(page_text.strip()),
            })
    else:
        all_text = []
        for item in content_list:
            text = _extract_text_from_item(item)
            if text:
                all_text.append(text)
        combined_text = "\n".join(all_text)
        pages.append({
            "page_num": 1,
            "text": combined_text,
            "char_count": len(combined_text),
        })

    return {
        "text": pages[0]["text"] if pages else "",
        "pages": pages,
        "total_pages": len(pages),
        "mineru_output_dir": mineru_output,
        "content_source": SOURCE_OCR,
        "ocr_status": OCR_PENDING,
    }


# ===== ZIP处理函数 =====

def process_zip(path, output_dir, parent_id=None):
    """解压ZIP并递归处理内部文件。

    对应Spec：ZIP（1个，马工程下册），解压后按内部文件类型路由处理。
    content_source和ocr_status取决于内部文件类型。

    Args:
        path: ZIP文件绝对路径
        output_dir: 输出目录
        parent_id: 父文件ID（用于标识嵌套处理）

    Returns:
        dict: 提取结果，包含extracted_files列表
    """
    # 创建临时解压目录
    temp_dir = tempfile.mkdtemp(prefix="wenyan_zip_")

    try:
        # 解压ZIP文件
        import zipfile
        with zipfile.ZipFile(path, "r") as zf:
            zf.extractall(temp_dir)

        # 遍历解压后的文件
        extracted_files = []
        for root, dirs, filenames in os.walk(temp_dir):
            for filename in filenames:
                if filename.startswith("~$") or filename == "Thumbs.db":
                    continue

                file_path = os.path.join(root, filename)
                rel_path = os.path.relpath(file_path, temp_dir)
                file_type = _get_file_type(filename)

                if file_type == "unknown":
                    print(f"  跳过ZIP内部不支持的文件: {rel_path}",
                          file=sys.stderr)
                    continue

                # 递归处理内部文件
                try:
                    inner_result = _route_by_file_type(
                        file_path, file_type, output_dir
                    )
                    extracted_files.append({
                        "file_name": filename,
                        "relative_path": rel_path.replace("\\", "/"),
                        "file_type": file_type,
                        "status": "completed",
                    })
                except Exception as e:
                    extracted_files.append({
                        "file_name": filename,
                        "relative_path": rel_path.replace("\\", "/"),
                        "file_type": file_type,
                        "status": "failed",
                        "error": str(e),
                    })

        return {
            "extracted_files": extracted_files,
            "total_extracted": len(extracted_files),
            "content_source": SOURCE_NATIVE,
            "ocr_status": OCR_VERIFIED,
        }

    finally:
        # 清理临时解压目录
        try:
            shutil.rmtree(temp_dir, ignore_errors=True)
        except Exception:
            pass


def _get_file_type(file_name):
    """根据扩展名判断文件类型（与scan_files.py中的逻辑一致）。"""
    ext_map = {
        ".pdf": "pdf", ".docx": "docx", ".doc": "doc",
        ".xlsx": "xlsx", ".xls": "xls",
        ".jpg": "image", ".jpeg": "image", ".png": "image",
        ".bmp": "image", ".tiff": "image", ".gif": "image",
        ".zip": "zip",
    }
    _, ext = os.path.splitext(file_name)
    return ext_map.get(ext.lower(), "unknown")


def _route_by_file_type(file_path, file_type, output_dir):
    """根据文件类型路由到对应处理函数（内部使用，不含重试逻辑）。

    Args:
        file_path: 文件绝对路径
        file_type: 文件类型字符串
        output_dir: 输出目录

    Returns:
        dict: 处理结果
    """
    if file_type == "pdf":
        return _route_pdf(file_path, output_dir)
    elif file_type == "docx":
        return process_docx(file_path)
    elif file_type == "doc":
        return process_doc(file_path)
    elif file_type == "xlsx":
        return process_xlsx(file_path)
    elif file_type == "xls":
        return process_xls(file_path)
    elif file_type == "image":
        return process_image(file_path)
    elif file_type == "zip":
        return process_zip(file_path, output_dir)
    else:
        raise ValueError(f"不支持的文件类型: {file_type}")


def _route_pdf(file_path, output_dir):
    """根据PDF类型路由到对应的PDF处理函数。

    Args:
        file_path: PDF文件绝对路径
        output_dir: 输出目录（用于创建临时工作目录）

    Returns:
        dict: 处理结果
    """
    # 创建临时工作目录用于MinerU输出
    work_dir = tempfile.mkdtemp(prefix="wenyan_pdf_")

    try:
        # 需要从manifest记录中获取pdf_type，这里通过重新检测
        # 在route_and_process中会传入pdf_type，这里作为后备
        pdf_type, _ = _quick_detect_pdf_type(file_path)

        if pdf_type == "native":
            return process_pdf_native(file_path)
        elif pdf_type == "ocr_layer":
            return process_pdf_ocr_layer(file_path)
        elif pdf_type == "scan_only":
            return process_pdf_scan_only(file_path, work_dir)
        elif pdf_type == "mixed":
            return process_pdf_mixed(file_path, work_dir)
        else:
            # 未知类型，默认用MinerU处理
            return process_pdf_scan_only(file_path, work_dir)
    finally:
        # 注意：不立即清理work_dir，因为MinerU输出可能需要后续查看
        # 实际清理在处理完成后由调用方或手动进行
        pass


def _quick_detect_pdf_type(pdf_path):
    """快速检测PDF类型（与scan_files.py中detect_pdf_type逻辑一致）。

    用于pipeline_runner在没有manifest中pdf_type信息时的后备检测。

    Args:
        pdf_path: PDF文件路径

    Returns:
        tuple: (pdf_type, needs_ocr)
    """
    try:
        import pdfplumber
    except ImportError:
        return "scan_only", True

    try:
        with pdfplumber.open(pdf_path) as pdf:
            if len(pdf.pages) == 0:
                return "scan_only", True

            page_char_counts = []
            for page in pdf.pages:
                text = page.extract_text() or ""
                page_char_counts.append(len(text.strip()))

            total_chars = sum(page_char_counts)
            total_pages = len(page_char_counts)
            has_text_pages = sum(1 for c in page_char_counts if c > 0)
            no_text_pages = total_pages - has_text_pages

            if has_text_pages == 0:
                return "scan_only", True
            if no_text_pages > 0:
                return "mixed", True

            avg_chars = total_chars / total_pages
            if avg_chars > 100:
                return "native", False
            else:
                return "ocr_layer", True
    except Exception:
        return "scan_only", True


# ===== 路由与重试 =====

def route_and_process(file_record, output_dir):
    """根据file_type和pdf_type路由到对应的处理函数。

    Args:
        file_record: manifest中的文件记录字典
        output_dir: 输出目录

    Returns:
        dict: 处理结果（包含content_source、ocr_status等）
    """
    file_type = file_record["file_type"]
    file_path = file_record["absolute_path"]

    if file_type == "pdf":
        pdf_type = file_record.get("pdf_type", "scan_only")
        work_dir = tempfile.mkdtemp(prefix="wenyan_pdf_")

        try:
            if pdf_type == "native":
                return process_pdf_native(file_path)
            elif pdf_type == "ocr_layer":
                return process_pdf_ocr_layer(file_path)
            elif pdf_type == "scan_only":
                return process_pdf_scan_only(file_path, work_dir)
            elif pdf_type == "mixed":
                return process_pdf_mixed(file_path, work_dir)
            else:
                # 未知PDF类型，保守用MinerU处理
                return process_pdf_scan_only(file_path, work_dir)
        finally:
            pass  # work_dir保留供后续查看MinerU输出

    elif file_type == "docx":
        return process_docx(file_path)

    elif file_type == "doc":
        return process_doc(file_path)

    elif file_type == "xlsx":
        return process_xlsx(file_path)

    elif file_type == "xls":
        return process_xls(file_path)

    elif file_type == "image":
        return process_image(file_path)

    elif file_type == "zip":
        return process_zip(file_path, output_dir,
                           parent_id=file_record.get("id"))

    else:
        raise ValueError(f"不支持的文件类型: {file_type}")


def process_with_retry(file_record, output_dir, max_retries=MAX_RETRIES):
    """带指数退避重试的处理包装器。

    重试策略：
      - 最多重试max_retries次
      - 每次重试前等待 BASE_BACKOFF * 2^(attempt-1) 秒
      - 所有重试失败后返回错误信息

    Args:
        file_record: 文件记录字典
        output_dir: 输出目录
        max_retries: 最大重试次数

    Returns:
        tuple: (success: bool, result: dict or error_message: str)
    """
    last_error = ""

    for attempt in range(1, max_retries + 1):
        try:
            result = route_and_process(file_record, output_dir)
            return True, result

        except Exception as e:
            last_error = str(e)
            print(f"  尝试 {attempt}/{max_retries} 失败: {last_error}",
                  file=sys.stderr)

            if attempt < max_retries:
                # 指数退避等待
                wait_time = BASE_BACKOFF_SECONDS * (2 ** (attempt - 1))
                print(f"  等待 {wait_time} 秒后重试...", file=sys.stderr)
                time.sleep(wait_time)

    return False, last_error


# ===== Manifest管理 =====

def load_manifest(manifest_path):
    """加载manifest.json文件。

    Args:
        manifest_path: manifest.json文件路径

    Returns:
        dict: manifest字典
    """
    with open(manifest_path, "r", encoding="utf-8") as f:
        return json.load(f)


def save_manifest(manifest, manifest_path):
    """保存manifest.json文件（断点续传关键操作）。

    Args:
        manifest: manifest字典
        manifest_path: manifest.json文件路径
    """
    # 先写入临时文件，再原子替换，防止写入中断导致文件损坏
    tmp_path = manifest_path + ".tmp"
    with open(tmp_path, "w", encoding="utf-8") as f:
        json.dump(manifest, f, ensure_ascii=False, indent=2)
    os.replace(tmp_path, manifest_path)


def update_file_status(manifest, file_id, status, error=None,
                       result_summary=None):
    """更新manifest中指定文件的状态。

    Args:
        manifest: manifest字典
        file_id: 文件ID
        status: 新状态（"completed"/"failed"/"pending"）
        error: 失败原因（仅status="failed"时）
        result_summary: 处理结果摘要（仅status="completed"时）
    """
    for file_record in manifest["files"]:
        if file_record["id"] == file_id:
            file_record["status"] = status
            if status == "failed" and error:
                file_record["error"] = error
            if status == "completed":
                if result_summary:
                    file_record["result_summary"] = result_summary
                # 清除可能残留的error字段（重试成功后stale error）
                file_record.pop("error", None)
            break


# ===== 主流程 =====

def _process_single_file(file_record, output_dir, manifest, manifest_path,
                         manifest_lock, index, total):
    """处理单个文件（线程安全，供并行/串行调用）。

    Args:
        file_record: manifest中的文件记录字典
        output_dir: 输出目录
        manifest: manifest字典（共享引用，写操作需加锁）
        manifest_path: manifest.json文件路径
        manifest_lock: 线程锁（保护manifest写操作）
        index: 当前文件序号（1-based，用于显示）
        total: 总文件数

    Returns:
        str: "success" / "failed" / "skipped"
    """
    file_id = file_record["id"]
    file_name = file_record["file_name"]
    file_type = file_record["file_type"]

    print(f"\n[{index}/{total}] 处理: {file_name}")
    print(f"  ID: {file_id}, 类型: {file_type}", end="")
    if file_type == "pdf":
        print(f", PDF类型: {file_record.get('pdf_type', 'unknown')}")
    else:
        print()

    # 跳过重复文件
    if file_record.get("is_duplicate", False):
        print(f"  跳过重复文件（原版: {file_record.get('duplicate_of')}）")
        with manifest_lock:
            file_record["status"] = "skipped"
            save_manifest(manifest, manifest_path)
        return "skipped"

    # 带重试的处理
    success, result = process_with_retry(file_record, output_dir)

    with manifest_lock:
        if success:
            # 处理成功：写入输出JSON，更新manifest
            output_json = {
                "id": file_id,
                "relative_path": file_record["relative_path"],
                "file_name": file_name,
                "file_type": file_type,
                "pdf_type": file_record.get("pdf_type"),
                "category": file_record.get("category"),
                "content_source": result.get("content_source"),
                "ocr_status": result.get("ocr_status"),
                "data": result,
                "processed_at": datetime.now().isoformat(),
                "status": "completed",
            }

            # 移除data中已提升到顶层的字段
            data_copy = dict(result)
            data_copy.pop("content_source", None)
            data_copy.pop("ocr_status", None)
            output_json["data"] = data_copy

            output_path = os.path.join(output_dir, f"{file_id}.json")
            with open(output_path, "w", encoding="utf-8") as f:
                json.dump(output_json, f, ensure_ascii=False, indent=2)

            update_file_status(
                manifest, file_id, "completed",
                result_summary={
                    "output_file": f"{file_id}.json",
                    "content_source": result.get("content_source"),
                    "ocr_status": result.get("ocr_status"),
                }
            )
            file_record["attempts"] = MAX_RETRIES
            print(f"  处理成功 → {output_path}")
            save_manifest(manifest, manifest_path)
            return "success"
        else:
            update_file_status(manifest, file_id, "failed", error=result)
            file_record["attempts"] = MAX_RETRIES
            print(f"  处理失败（已重试{MAX_RETRIES}次）: {result}",
                  file=sys.stderr)
            save_manifest(manifest, manifest_path)
            return "failed"


def run_pipeline(manifest_path, output_dir, resume=False, workers=DEFAULT_FAST_WORKERS):
    """运行批处理管线。

    优化（2026-07-11）：
      - 按文件类型优先级排序（快速文件优先）
      - 快速文件（docx/xlsx/native等）并行处理
      - 慢速文件（scan_only/mixed/image）串行处理（MinerU内存占用大）
      - mixed PDF只对扫描页调用MinerU

    Args:
        manifest_path: manifest.json路径
        output_dir: 输出目录
        resume: 断点续传模式（跳过已失败的文件）
        workers: 快速文件并行工作线程数
    """
    # 加载manifest
    manifest = load_manifest(manifest_path)
    print(f"已加载manifest: {manifest_path}")
    print(f"文件总数: {manifest['total']}")

    # 确保输出目录存在
    os.makedirs(output_dir, exist_ok=True)

    # 确定待处理文件列表
    if not resume:
        reset_count = 0
        for file_record in manifest["files"]:
            if file_record["status"] == "failed":
                file_record["status"] = "pending"
                file_record["attempts"] = 0
                reset_count += 1
        if reset_count > 0:
            print(f"非断点续传模式：重置 {reset_count} 个失败文件为待处理")
            save_manifest(manifest, manifest_path)

    pending_files = [f for f in manifest["files"]
                     if f["status"] == "pending"]
    print(f"待处理文件: {len(pending_files)} 个")

    # 按优先级排序：快速文件优先，同优先级按页数升序
    pending_files.sort(
        key=lambda f: (get_file_priority(f), f.get("page_count", 0) or 0)
    )

    # 分两类：快速文件（并行）和慢速文件（串行，MinerU内存大不能并行）
    fast_files = [f for f in pending_files if get_file_priority(f) <= 2]
    slow_files = [f for f in pending_files if get_file_priority(f) > 2]

    print(f"  快速文件（并行处理, workers={workers}）: {len(fast_files)} 个")
    print(f"  慢速文件（串行OCR）: {len(slow_files)} 个")

    success_count = 0
    fail_count = 0
    skip_count = 0
    manifest_lock = threading.Lock()

    # 阶段1：快速文件并行处理
    if fast_files:
        print(f"\n{'=' * 50}")
        print(f"阶段1：并行处理 {len(fast_files)} 个快速文件")
        print(f"{'=' * 50}")

        with ThreadPoolExecutor(max_workers=min(workers, len(fast_files))) as executor:
            futures = {}
            for i, file_record in enumerate(fast_files):
                future = executor.submit(
                    _process_single_file,
                    file_record, output_dir, manifest, manifest_path,
                    manifest_lock, i + 1, len(fast_files)
                )
                futures[future] = file_record

            for future in as_completed(futures):
                status = future.result()
                if status == "success":
                    success_count += 1
                elif status == "skipped":
                    skip_count += 1
                else:
                    fail_count += 1

    # 阶段2：慢速文件串行处理（MinerU内存占用大，不能并行）
    if slow_files:
        print(f"\n{'=' * 50}")
        print(f"阶段2：串行处理 {len(slow_files)} 个慢速文件（MinerU OCR）")
        print(f"{'=' * 50}")

        for i, file_record in enumerate(slow_files):
            status = _process_single_file(
                file_record, output_dir, manifest, manifest_path,
                manifest_lock, i + 1, len(slow_files)
            )
            if status == "success":
                success_count += 1
            elif status == "skipped":
                skip_count += 1
            else:
                fail_count += 1

    # 打印汇总
    print("\n" + "=" * 50)
    print("处理完成汇总:")
    print(f"  成功: {success_count}")
    print(f"  失败: {fail_count}")
    print(f"  跳过(重复): {skip_count}")
    print(f"  总计: {len(pending_files)}")
    print("=" * 50)

    # 生成失败文件清单
    if fail_count > 0:
        failed_list = [
            f for f in manifest["files"] if f["status"] == "failed"
        ]
        failed_report_path = os.path.join(output_dir, "failed_files.json")
        with open(failed_report_path, "w", encoding="utf-8") as f:
            json.dump(
                [{"id": f["id"], "file_name": f["file_name"],
                  "relative_path": f["relative_path"],
                  "error": f.get("error", "unknown")}
                 for f in failed_list],
                f, ensure_ascii=False, indent=2
            )
        print(f"失败文件清单: {failed_report_path}")


# ===== 命令行入口 =====

def main():
    """命令行入口函数。"""
    parser = argparse.ArgumentParser(
        description="文研App资料数字化处理管线。"
                    "manifest驱动的断点续传批处理脚本。",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
  # 首次运行（处理所有pending文件，重置failed文件）
  python pipeline_runner.py --input manifest.json --output output

  # 断点续传（跳过已completed和failed的文件）
  python pipeline_runner.py --input manifest.json --output output --resume

处理路由:
  pdf/native    → pdfplumber直接提取
  pdf/ocr_layer → pdfplumber提取OCR文本层
  pdf/scan_only → MinerU CLI (mineru -p <input> -o <output> -m auto)
  pdf/mixed     → pdfplumber + MinerU混合处理
  docx          → python-docx提取
  doc           → pywin32(win32com)转docx后提取
  xlsx          → openpyxl提取
  xls           → xlrd 2.0+提取
  image         → MinerU CLI pipeline后端OCR（2026-07-10因PaddleOCR 3.x不兼容改用MinerU）
  zip           → 解压后递归处理
        """,
    )
    parser.add_argument(
        "--input",
        default=None,
        help="manifest.json路径（默认: <脚本目录>/manifest.json）",
    )
    parser.add_argument(
        "--output",
        default=None,
        help="输出目录路径（默认: <脚本目录>/output）",
    )
    parser.add_argument(
        "--resume",
        action="store_true",
        help="断点续传模式（跳过已completed和failed的文件，"
             "不重置failed文件为pending）",
    )
    parser.add_argument(
        "--workers",
        type=int,
        default=DEFAULT_FAST_WORKERS,
        help=f"快速文件并行工作线程数（默认{DEFAULT_FAST_WORKERS}，"
             f"慢速MinerU OCR文件始终串行）",
    )

    args = parser.parse_args()

    # 确定路径
    script_dir = os.path.dirname(os.path.abspath(__file__))
    manifest_path = args.input or os.path.join(script_dir, "manifest.json")
    output_dir = args.output or os.path.join(script_dir, "output")

    # 验证manifest文件存在
    if not os.path.exists(manifest_path):
        print(f"错误：manifest文件不存在: {manifest_path}", file=sys.stderr)
        print(f"请先运行: python {os.path.join(script_dir, 'scan_files.py')}")
        sys.exit(1)

    # 运行管线
    run_pipeline(manifest_path, output_dir, resume=args.resume,
                 workers=args.workers)


if __name__ == "__main__":
    main()
